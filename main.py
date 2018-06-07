from urllib.request import urlopen
from urllib.request import urlretrieve 
from urllib.error import HTTPError, URLError
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import os
import re 
import time 
#打印今天的日期
today = time.strftime("%Y-%m-%d", time.localtime())
print(today)
#将今天的日期分解，用于后期对当日文章的正则匹配
tm_year = str(time.localtime()[0])
tm_mon = str(time.localtime()[1])
tm_day = str(time.localtime()[2])
#常速英语的网址，也可以是慢速英语的网址
url = "http://www.51voa.com/VOA_Standard_English/"

def init():
    '''
    初始化函数，用于生成存放文章的文件夹，以当日日期命名
    '''
    path = "./" + today 
    folder = os.path.exists(path)
    if not folder:                   
        os.makedirs(path)
        print("ok")
    else:
        print("exited")

def get_html_bsobj(url):
    '''
    获得单个网页的bsobj对象
    '''
    try:
        html = urlopen(url)
    except (HTTPError, URLError) as e:
        return None
    bsObj = BeautifulSoup(html.read(), 'html.parser')
    return bsObj 

def get_articles_list(bsObj):
    '''
    通过正则表达式，在列表网页对象中筛选出今天的所有文章链接
    '''
    pattern = ".*" + tm_year + "-" + tm_mon + "-" + tm_day + ".*"
    print(pattern)
    sites = bsObj.select('div[id="list"] > ul > li > a')
    #在网页中筛选出所有的链接放入sites列表
    links = [] 
    for site in sites:
        if re.match(pattern , site.get_text()) != None:
            links.append("http://www.51voa.com" + site.get("href"))
    #从列表中获取所有的文章链接进入列表，并返回这个列表
    return links 
    
def get_article(article_url):
    '''
    通过文章的链接获取单个文章，并保存到一个docx文档
    '''
    article_bsobj = get_html_bsobj(article_url)
    #获得文章的bsobj对象
    #下面是通过对网页HTML代码的分析筛选出内容
    title = article_bsobj.select('div[id="title"]')[0].get_text()
    #获取文章的标题
    author = article_bsobj.select('div[id="content"] > span[class="byline"]')[0].get_text()
    #获取文章的作者
    time = article_bsobj.select('div[id="content"] > span[class="datetime"]')[0].get_text()
    #获取文章的发布时间
    article = article_bsobj.select('div[id="content"] > p')
    #获取文章的正文
    document = Document()
    document.add_heading(time, 0)
    document.add_heading(title, level=1)
    document.add_heading(author, level=3)
    for each in article:
        document.add_paragraph(each.get_text())
    #保存到docx文档，具体方法建议查看pydocx的帮助文档
    doc_name = "./" + today + "/" + title + ".docx"
    print(doc_name)
    document.save(doc_name)
    
    mp3_link = article_bsobj.select('div[id="menubar"] > a[id="mp3"]')[0].get('href')
    mp3_name = "./" + today + "/" + title + ".mp3"
    print(mp3_name)
    urlretrieve(mp3_link, mp3_name)
    #下载MP3

if __name__ == '__main__':
    init()
    article_lists = get_articles_list(get_html_bsobj(url))
    for each_article in article_lists:
        get_article(each_article)
